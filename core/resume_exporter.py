"""
Smart Resume — Resume Exporter v3
====================================
Pure Python — NO Node.js required.
  - DOCX via python-docx  (already in requirements.txt)
  - PDF  via reportlab    (pip install reportlab)

ALL content parsed from LLM-generated text.
"""
from __future__ import annotations
import logging, os, re
from io import BytesIO
from typing import Optional

logger = logging.getLogger("smart_resume.exporter")

ALL_SECTIONS = ["SUMMARY","OBJECTIVE","TECHNICAL SKILLS","SKILLS","EXPERIENCE","WORK EXPERIENCE","PROJECTS","EDUCATION","CERTIFICATIONS"]
ACCENT = (0, 0, 0)

def _clean_email(email: str) -> str:
    if not email: return ""
    email = email.strip().lower()
    providers = ["gmail", "outlook", "hotmail", "yahoo", "icloud"]
    for p in providers:
        if f"@{p}.co" in email:
            return email.replace(f"@{p}.co", f"@{p}.com")
    return email


# ─────────────────────────────────────────────────────────────────────────────
# PUBLIC API
# ─────────────────────────────────────────────────────────────────────────────

def export_resume(resume_text: str, parsed_resume, fmt: str = "docx") -> bytes:
    d = _build_resume_data(resume_text, parsed_resume)
    if fmt == "docx": return _export_docx(d)
    if fmt == "pdf":  return _export_pdf(d)
    raise ValueError(f"Unknown format: {fmt}")


# ─────────────────────────────────────────────────────────────────────────────
# DATA BUILDER — parses plain text resume (preview = what user sees)
# ─────────────────────────────────────────────────────────────────────────────

# All known section header names — used to detect section boundaries
_SECTION_HEADERS = re.compile(
    r"^(OBJECTIVE|SUMMARY|PROFILE|SKILLS|TECHNICAL SKILLS|WORK EXPERIENCE|EXPERIENCE|"
    r"PROJECTS|EDUCATION|CERTIFICATIONS|ACHIEVEMENTS|CERTIFICATIONS & ACHIEVEMENTS)\s*$",
    re.IGNORECASE | re.MULTILINE,
)

def _extract_section(text: str, name: str) -> str:
    """Extract content of a named section from plain text resume."""
    # Find start of this section
    start_m = re.search(
        r"(?:^|\n)" + re.escape(name) + r"(?::)?\s*\n",
        text, re.IGNORECASE
    )
    if not start_m:
        return ""
    start = start_m.end()

    # Find end = next section header
    end_m = _SECTION_HEADERS.search(text, start)
    end   = end_m.start() if end_m else len(text)

    return text[start:end].strip()


def _build_resume_data(resume_text: str, pr) -> dict:
    """
    Build resume data dict from plain text resume (Groq output).
    This ensures DOCX/PDF matches the Preview exactly.
    Falls back to ParsedResume fields only for contact info.
    """
    # ── Contact from ParsedResume (more reliable than parsing plain text) ─────
    name     = pr.name or "Candidate"
    linkedin = getattr(pr, "linkedin", None) or ""
    github   = getattr(pr, "github",   None) or ""
    email    = _clean_email(pr.email or "")
    cp       = [p for p in [pr.phone or "", email, pr.location or ""] if p.strip()]
    if linkedin: cp.append(linkedin)
    if github:   cp.append(github)
    contact  = "  |  ".join(cp)

    # ── All sections from plain text resume ──────────────────────────────────
    objective = (
        _extract_section(resume_text, "OBJECTIVE") or
        _extract_section(resume_text, "SUMMARY") or
        _extract_section(resume_text, "PROFILE") or
        pr.summary or ""
    )

    # Skills — parse from plain text
    sb = (
        _extract_section(resume_text, "SKILLS") or
        _extract_section(resume_text, "TECHNICAL SKILLS")
    )
    skills = {}
    if sb:
        for line in sb.splitlines():
            line = line.strip().lstrip("-•* ")
            if not line:
                continue
            if ":" in line:
                cat, _, vals = line.partition(":")
                items = [v.strip() for v in vals.split(",") if v.strip()]
                if items:
                    skills[cat.strip()] = items
    if not skills:
        skills = _auto_categorize(pr.skills or [])
    skills = {k: v for k, v in skills.items() if v}

    # Experience — parse from plain text
    eb = (
        _extract_section(resume_text, "WORK EXPERIENCE") or
        _extract_section(resume_text, "EXPERIENCE")
    )
    # Filter out empty experience sections (just blank lines or dashes)
    if eb and re.match(r"^[-\s]*$", eb):
        eb = ""
    experience = _parse_exp(eb) if eb else []
    # Fallback to ParsedResume only if plain text has nothing
    if not experience and pr.experience:
        experience = [{"role": e.role, "company": e.company,
                       "duration": e.duration or "",
                       "responsibilities": e.responsibilities}
                      for e in pr.experience]

    # Projects — structured data from ParsedResume (LLM quality), fallback to text
    projects = []
    if getattr(pr, "projects", None):
        for p in pr.projects:
            projects.append({
                "title":        getattr(p, "title", "Project") or "Project",
                "description":  getattr(p, "description", "") or "",
                "technologies": getattr(p, "technologies", []) or [],
                "metrics":      getattr(p, "metrics", []) or [],
            })
    if not projects:
        pb = _extract_section(resume_text, "PROJECTS")
        projects = _parse_proj(pb) if pb else []

    # Education — structured data from LLM (clean names), fallback to text parse
    education = []
    if getattr(pr, "education", None):
        from datetime import datetime as _dt
        _cur_yr = _dt.now().year
        for e in pr.education:
            inst     = getattr(e, "institution", "") or ""
            deg      = getattr(e, "degree", "") or ""
            field    = getattr(e, "field", "") or getattr(e, "branch", "") or ""
            grad_val = getattr(e, "graduation_year", "") or ""
            gpa_val  = getattr(e, "gpa", "") or ""
            level    = getattr(e, "level", "").lower() if hasattr(e, "level") else ""
            # Format graduation year — preserve ranges as-is
            grad_str = str(grad_val).strip()
            # Only do year-arithmetic if it looks like a single 4-digit year
            _single_yr = re.match(r'^\d{4}$', grad_str)
            if _single_yr:
                try:
                    yr_int = int(grad_str)
                    grad_str = f"Expected {yr_int}" if yr_int > _cur_yr else (
                        f"2023 - Present" if yr_int == _cur_yr else grad_str
                    )
                except Exception:
                    pass  # keep original string
            # Compose left side
            course = f"{deg} in {field}" if field else deg
            education.append({
                "institution":     inst,
                "degree":          course,
                "graduation_year": grad_str,
                "gpa":             gpa_val,
                "level":           level,
            })
    if not education:
        education = _parse_education(resume_text, pr)

    # Certifications — parse from plain text
    cb = _extract_section(resume_text, "CERTIFICATIONS")
    certifications = []
    if cb:
        for line in cb.splitlines():
            line = line.strip().lstrip("-•* ")
            # Skip placeholder lines
            if not line or line == "-" or re.match(r"^-+$", line):
                continue
            if len(line) > 3:
                ym = re.search(r"(20\d\d)", line)
                certifications.append({
                    "name":   line,
                    "issuer": "",
                    "year":   int(ym.group(1)) if ym else "",
                })
    if not certifications:
        certifications = [{"name": c.name, "issuer": c.issuer or "", "year": c.year or ""}
                          for c in pr.certifications]

    # Redistribute competitions from experience to projects
    final_exp  = []
    extra_proj = []
    comp_kw    = {"competition","hackathon","abaja","contest","baja","robocon"}
    for exp in experience:
        is_comp = any(kw in (exp.get("role","") + exp.get("company","")).lower()
                      for kw in comp_kw)
        if is_comp:
            # Smart Title: prefer company if role is generic (e.g., Member | A BAJA)
            role_l = exp.get("role", "").lower()
            comp_l = exp.get("company", "").lower()
            generic = {"member", "lead", "developer", "participant", "student", "intern", "engineer", "associate"}
            if any(r in role_l for r in generic) and exp.get("company"):
                title = exp.get("company")
            else:
                title = exp.get("role") or exp.get("company") or "Competition"

            extra_proj.append({
                "title":        title,
                "description":  f"Role: {exp.get('role', '')}",
                "technologies": [],
                "metrics":      exp.get("responsibilities", []),
            })
        else:
            final_exp.append(exp)

    return {
        "name":           name,
        "contact":        contact,
        "objective":      objective,
        "skills":         skills,
        "experience":     final_exp,
        "projects":       projects + extra_proj,
        "education":      education,
        "certifications": certifications,
        "school_10th":    _parse_school(resume_text, "10"),
        "school_12th":    _parse_school(resume_text, "12"),
    }


def _parse_education(resume_text: str, pr) -> list:
    """
    Parse education from plain text resume.
    Handles:
      Bachelor of Technology, MA College of Engineering (2023) | CGPA: 7.79
      Class XII: G HHS MEZHATHUR (2022) | 96
    Falls back to ParsedResume education entries.
    """
    from datetime import datetime as _dte
    _cur_yr = _dte.now().year

    eb = _extract_section(resume_text, "EDUCATION")
    if not eb:
        # Fallback
        def _yr(e):
            v = str(e.graduation_year).strip() if e.graduation_year else ""
            if not v: return ""
            # If it already looks like a range, keep it verbatim
            if re.search(r"20\d\d\s*[-\u2013\u2014to]\s*(20\d\d|[Pp]resent)", v):
                return v
            try:
                yr = int(re.match(r'^(\d{4})', v).group(1))
                if yr > _cur_yr: return f"Expected {yr}"
                return str(yr)
            except Exception:
                return v
        return [{"degree": e.degree, "institution": e.institution,
                 "graduation_year": _yr(e),
                 "gpa": f"CGPA: {e.gpa}" if e.gpa else ""}
                for e in pr.education]

    education = []
    for line in eb.splitlines():
        line = line.strip()
        if not line or re.match(r"^[-\s]*$", line):
            continue
        # Skip Class X/XII lines (handled via school_10th/school_12th)
        if re.match(r"(?i)class\s*(x|xii|10|12)", line):
            continue

        # Parse: "Degree, Institution (Year) | CGPA: X"
        gpa_m = re.search(r"\|\s*CGPA[:\s]*([\d.]+)", line, re.IGNORECASE)
        yr_m  = re.search(r"\((\d{4})\)", line)
        gpa   = f"CGPA: {gpa_m.group(1)}" if gpa_m else ""
        yr    = yr_m.group(1) if yr_m else ""

        # Remove year and gpa from main text
        main = re.sub(r"\s*\(\d{4}\)", "", line)
        main = re.sub(r"\s*\|.*$", "", main).strip()

        # Split degree and institution on comma
        if "," in main:
            parts = main.split(",", 1)
            degree = parts[0].strip()
            inst   = parts[1].strip()
        else:
            degree = main
            inst   = ""

        if degree and len(degree) > 2:
            education.append({
                "degree":          degree,
                "institution":     inst,
                "graduation_year": yr,
                "gpa":             gpa,
            })

    # If nothing parsed, fallback to ParsedResume
    if not education:
        def _yr2(e):
            v = str(e.graduation_year).strip() if e.graduation_year else ""
            if not v: return ""
            # If it already looks like a range, keep it verbatim
            if re.search(r"20\d\d\s*[-\u2013\u2014to]\s*(20\d\d|[Pp]resent)", v):
                return v
            try:
                yr = int(re.match(r'^(\d{4})', v).group(1))
                return f"Expected {yr}" if yr > _cur_yr else str(yr)
            except Exception:
                return v
        education = [{"degree": e.degree, "institution": e.institution,
                      "graduation_year": _yr2(e),
                      "gpa": f"CGPA: {e.gpa}" if e.gpa else ""}
                     for e in pr.education]

    return education



def _parse_school(resume_text: str, level: str) -> str | None:
    """
    Extract Class X or Class XII line from plain text resume.
    level: "10" or "12"
    Returns formatted string like "G HHS MEZHATHUR | 2022 | 96%" or None.
    Skips placeholder lines like "Class X: -".
    """
    if level == "12":
        pat = r"(?i)class\s*(?:xii?|12th?)[:\s]+(.*)"
    else:
        pat = r"(?i)class\s*(?:x\b|10th?)[:\s]+(.*)"

    m = re.search(pat, resume_text)
    if not m:
        return None

    val = m.group(1).strip()
    # Reject placeholder values
    if not val or val == "-" or re.match(r"^-+$", val):
        return None
    # Reject bare numbers (hallucinated)
    if re.match(r"^\d{1,3}$", val):
        return None

    return val


def _parse_exp(block: str) -> list:
    exps, cur = [], None
    for raw in block.splitlines():
        line = raw.strip()
        if not line: continue
        is_bul = line.startswith(("-","•","*"))
        if not is_bul and ("|" in line or re.match(r"^[A-Z][A-Za-z /]+(?:\||-|\d)", line)):
            if cur: exps.append(cur)
            parts = [p.strip() for p in line.split("|")]
            cur = {"role":parts[0],"company":parts[1] if len(parts)>1 else "","duration":parts[2] if len(parts)>2 else "","responsibilities":[]}
        elif is_bul:
            if cur is None: cur={"role":"","company":"","duration":"","responsibilities":[]}
            t = line.lstrip("-•* ").strip()
            if t: cur["responsibilities"].append(t)
    if cur: exps.append(cur)
    return [e for e in exps if e.get("role") or e.get("responsibilities")]

def _parse_proj(block: str) -> list:
    projs, cur = [], None
    for raw in block.splitlines():
        line = raw.strip()
        if not line: continue
        is_bul = line.startswith(("-","•","*"))
        if not is_bul:
            if cur: projs.append(cur)
            if "|" in line:
                parts  = [p.strip() for p in line.split("|",1)]
                title  = parts[0]
                tech_s = re.sub(r"(?i)^tech\s*:\s*","",parts[1]) if len(parts)>1 else ""
                # Split on comma but preserve known multi-word frameworks
                raw_techs = [t.strip() for t in tech_s.split(",") if t.strip()]
                # Merge consecutive tokens that form known frameworks
                MULTI_WORD = {"python flask","python django","python fastapi",
                              "react js","node js","next js","ruby rails",
                              "spring boot","asp net","machine learning","deep learning"}
                techs = []
                i = 0
                while i < len(raw_techs):
                    if i + 1 < len(raw_techs):
                        combined = (raw_techs[i] + " " + raw_techs[i+1]).lower()
                        if combined in MULTI_WORD:
                            techs.append(raw_techs[i] + " " + raw_techs[i+1])
                            i += 2
                            continue
                    techs.append(raw_techs[i])
                    i += 1
            else:
                title, techs = line, []
            cur = {"title":title,"description":"","technologies":techs,"metrics":[]}
        else:
            if cur is None: cur={"title":"","description":"","technologies":[],"metrics":[]}
            t = line.lstrip("-•* ").strip()
            if not cur["description"]: cur["description"] = t
            elif t: cur["metrics"].append(t)
    if cur: projs.append(cur)
    return [p for p in projs if p.get("title") or p.get("description")]

def _auto_categorize(skills_list: list) -> dict:
    """
    Categorize a flat skills list into named buckets.

    CAPITALISATION RULE: every skill gets Title Case unless it is in the
    _PRESERVE set (fixed-case acronyms) or contains dots/slashes (e.g. React.js).

    "Other" RULE: only genuine skill names land here.  Single-word lowercase
    fragments that look like process nouns or substrings (e.g. "ing", "solving",
    "mak", "think") are silently dropped — they are parser artefacts.
    """
    import re as _re2

    # ── Exact-match sets (all lowercase) ──────────────────────────────────────
    LANGS = {
        "python","java","javascript","typescript","c","c++","c#","go","rust",
        "kotlin","swift","r","scala","ruby","php","bash","sql","dart","matlab",
        "vhdl","verilog",
    }
    FWKS = {
        "react","vue","angular","nextjs","django","flask","fastapi","spring",
        "pytorch","tensorflow","keras","ros","ros2","express","flutter",
        "langchain","opencv","numpy","pandas","scikit","sklearn","scikit-learn",
        "gazebo","moveit","rviz","streamlit","scipy","matplotlib","seaborn",
        "huggingface","transformers",
    }
    DBS = {
        "postgresql","mysql","sqlite","mongodb","redis","firebase","cassandra",
        "dynamodb","supabase","elasticsearch","sql server","oracle",
    }
    TOOLS = {
        "docker","kubernetes","git","github","linux","aws","gcp","azure",
        "jenkins","terraform","grafana","postman","cmake","arduino","raspberry",
        "simulink","keil","proteus","labview","autocad","solidworks","figma",
        "jira","vscode","vs code","jupyter","jupyter notebook","jupyter lab",
        "gitlab","bitbucket","esp32","raspberry pi","heroku","vercel","netlify",
    }
    WEB = {
        "react.js","react","next.js","nextjs","tailwind css","tailwind",
        "css","css3","html","html5","web development","web","frontend",
        "bootstrap","sass","scss",
    }
    BACKEND = {
        "node.js","node","express.js","express","fastapi","flask","django",
        "firebase","supabase","rest api","rest apis","apis","backend",
        "graphql","grpc",
    }
    AI = {
        "machine learning","deep learning","neural networks","nlp",
        "computer vision","artificial intelligence","ai","ml","data science",
        "data visualization","agentic ai","agentic","prompt engineering",
        "feature engineering","data preprocessing","data analysis",
        "natural language processing","reinforcement learning",
    }
    CORE_CS = {
        "data structures","data structures & algorithms","data structures and algorithms",
        "algorithms","dsa","oop","object oriented programming","oops",
        "dbms","database management","database management systems",
        "operating systems","os","computer networks","networking",
        "computer organization","computer architecture","discrete mathematics",
        "theory of computation","compiler design","software engineering",
        "system design","design patterns",
    }
    CYBER = {
        "vapt","penetration testing","cybersecurity","network security",
        "kali linux","metasploit","wireshark","security auditing",
        "ethical hacking","vulnerability assessment","devops",
    }
    SOFT = {
        "communication","leadership","teamwork","problem solving",
        "time management","adaptability","collaboration","creativity",
        "critical thinking","interpersonal","presentation","negotiation",
        "decision making","self motivated","multitasking",
        "attention to detail","work ethic","flexibility","team management",
        "team managment",  # tolerate common misspelling
    }

    # ── Substring/keyword sets for partial matching ────────────────────────────
    _AI_KW    = list(AI)
    _CYBER_KW = list(CYBER)
    _BACK_KW  = list(BACKEND)
    _WEB_KW   = list(WEB)
    _SOFT_KW  = list(SOFT)
    _CORE_KW  = list(CORE_CS)

    # ── Capitalisation helper ──────────────────────────────────────────────────
    _PRESERVE = {
        "C++","C#","SQL","HTML","CSS","API","AI","ML","VS Code","CSS3",
        "HTML5","REST","APIs","UI/UX","OOP","OOPS","DSA","DBMS","OS",
        "IoT","ESP32","NLP","GCP","AWS","GIT",
    }
    _ACRONYM_UPPER = {p.upper() for p in _PRESERVE}

    def _cap(s: str) -> str:
        """Return the correctly-capitalised form of a skill string."""
        stripped = s.strip()
        if stripped.upper() in _ACRONYM_UPPER:
            # Return from _PRESERVE map so we get the right mixed case (e.g. 'VS Code')
            for p in _PRESERVE:
                if p.upper() == stripped.upper():
                    return p
        # Keep things like React.js, Node.js, scikit-learn as-is
        if _re2.search(r"[._-]", stripped):
            # Still capitalise the first letter only if it is lowercase
            return stripped[0].upper() + stripped[1:] if stripped else stripped
        return stripped.title() if stripped else stripped

    # ── Junk-word filter for "Other" ───────────────────────────────────────────
    # Single-token words that are clearly parser artefacts / process nouns /
    # English gerund fragments — never valid standalone skill names.
    _JUNK_SINGLE = {
        "ing","mak","think","solv","mak","works","work","net",
        "structures","structure","system","systems","code","notebook",
        "processing","engineering","development","management","analysis",
        "algorithms","algorithm","computer","data","feature","programming",
        "software","networks","network","operating","decision","problem",
        "critical","thinking","solving","making","team","skills","skill",
        "vs","and","of","in","for","with","the","to",
    }

    def _is_junk(s: str) -> bool:
        """Return True if s looks like a parser fragment, not a real skill."""
        sl = s.strip().lower()
        # Single word, 3 chars or fewer → drop
        if " " not in sl and len(sl) <= 3 and sl not in {"c","r","go","c#","ai","ml","os"}:
            return True
        # In the explicit junk set
        if sl in _JUNK_SINGLE:
            return True
        # Ends with a hyphen-broken suffix  e.g. "mak-" or "think-"
        if _re2.search(r"-$", sl):
            return True
        return False

    # ── Bucket definitions ─────────────────────────────────────────────────────
    cats = {
        "Languages":             [],
        "Web Development":       [],
        "AI & Data Science":     [],
        "Core CS":               [],
        "Frameworks & Libraries":[],
        "Databases":             [],
        "Cybersecurity & Tools": [],
        "Tools & Platforms":     [],
        "Soft Skills":           [],
        "Other":                 [],
    }

    for raw_s in skills_list:
        sl     = raw_s.lower().strip()
        s_out  = _cap(raw_s.strip())

        if sl in LANGS:
            cats["Languages"].append(s_out)
        elif any(k in sl for k in _AI_KW):
            cats["AI & Data Science"].append(s_out)
        elif any(k in sl for k in _CORE_KW):
            cats["Core CS"].append(s_out)
        elif any(k in sl for k in _CYBER_KW):
            cats["Cybersecurity & Tools"].append(s_out)
        elif any(k in sl for k in _BACK_KW):
            cats["Frameworks & Libraries"].append(s_out)  # backend fwks
        elif any(k in sl for k in _WEB_KW):
            cats["Web Development"].append(s_out)
        elif sl in FWKS:
            cats["Frameworks & Libraries"].append(s_out)
        elif sl in DBS:
            cats["Databases"].append(s_out)
        elif sl in TOOLS:
            cats["Tools & Platforms"].append(s_out)
        elif any(soft in sl for soft in _SOFT_KW):
            cats["Soft Skills"].append(s_out)
        else:
            # Only add to Other if it looks like a genuine skill name
            if not _is_junk(raw_s):
                cats["Other"].append(s_out)

    return {k: v for k, v in cats.items() if v}


# ─────────────────────────────────────────────────────────────────────────────
# DOCX EXPORT
# ─────────────────────────────────────────────────────────────────────────────

def _export_docx(d: dict) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Cm(1.27)
        sec.bottom_margin = Cm(1.27)
        sec.left_margin   = Cm(1.27)
        sec.right_margin  = Cm(1.27)

    DARK = RGBColor(0, 0, 0)
    GREY = RGBColor(60, 60, 60)

    def _set_run(run, size, bold=False, color=None, italic=False, font="Times New Roman"):
        run.bold      = bold
        run.italic    = italic
        run.font.size = Pt(size)
        run.font.name = font
        run.font.color.rgb = color or DARK

    def _add_hr(para):
        pPr    = para._p.get_or_add_pPr()
        pBdr   = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _section_heading(text):
        p   = doc.add_paragraph()
        run = p.add_run(text.upper()[:1] + text.upper()[1:].lower())
        run.font.small_caps = True
        _set_run(run, 13, color=DARK)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        _add_hr(p)
        return p

    def _bullet(text):
        p   = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        _set_run(run, 11)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Pt(12)
        return p

    # Name
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = name_p.add_run(d.get("name",""))
    _set_run(nr, 24, color=DARK)
    name_p.paragraph_format.space_after = Pt(2)

    # Contact
    if d.get("contact"):
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(d["contact"])
        _set_run(cr, 10, color=DARK)
        cp.paragraph_format.space_after = Pt(8)

    # Summary
    obj = d.get("objective","")
    if obj and obj.strip():
        _section_heading("Summary")
        op = doc.add_paragraph()
        op.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        or_ = op.add_run(obj)
        _set_run(or_, 11)
        op.paragraph_format.space_after = Pt(4)

    # Work Experience
    experience = d.get("experience", [])
    if experience:
        _section_heading("Experience")
        for exp in experience:
            p         = doc.add_paragraph()
            left_text = exp.get("role","")
            if exp.get("company"): left_text += f" — {exp['company']}"
            rr = p.add_run(left_text)
            _set_run(rr, 11, bold=True)
            if exp.get("duration"):
                from docx.enum.text import WD_TAB_ALIGNMENT
                p.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)
                dr = p.add_run(f"\t{exp['duration']}")
                _set_run(dr, 11, italic=True, color=GREY)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)
            for resp in exp.get("responsibilities", []):
                if resp and resp.strip(): _bullet(resp.strip())

    # Projects
    projects = d.get("projects", [])
    if projects:
        _section_heading("Projects")
        for proj in projects:
            p  = doc.add_paragraph()
            tr = p.add_run(proj.get("title",""))
            _set_run(tr, 11, bold=True)
            techs = proj.get("technologies", [])
            if techs:
                techstr = p.add_run("  |  " + ", ".join(techs))
                _set_run(techstr, 11, italic=True)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)
            if proj.get("description","").strip(): _bullet(proj["description"].strip())
            for m in proj.get("metrics", []):
                if m and m.strip(): _bullet(m.strip())

    # Education
    education = d.get("education", [])
    if education:
        _section_heading("Education")
        for edu in education:
            p         = doc.add_paragraph()
            left_text = edu.get("degree","")
            if edu.get("institution"): left_text += f", {edu['institution']}"
            mr = p.add_run(left_text)
            _set_run(mr, 11)
            right_parts = []
            if edu.get("graduation_year"): right_parts.append(str(edu["graduation_year"]))
            if edu.get("gpa"):             right_parts.append(edu["gpa"])
            if right_parts:
                from docx.enum.text import WD_TAB_ALIGNMENT
                p.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)
                sr = p.add_run(f"\t{' | '.join(right_parts)}")
                _set_run(sr, 11)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(1)

        s10 = d.get("school_10th")
        s12 = d.get("school_12th")
        if s10 or s12:
            p     = doc.add_paragraph()
            parts = []
            if s12: parts.append(f"Class XII: {s12}")
            if s10: parts.append(f"Class X: {s10}")
            tr = p.add_run(" | ".join(parts))
            _set_run(tr, 10, italic=True, color=GREY)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)

    # Skills
    skills = d.get("skills", {})
    if skills:
        _section_heading("Skills")
        for cat, items in skills.items():
            if not items: continue
            p  = doc.add_paragraph()
            br = p.add_run(f"{cat}: ")
            _set_run(br, 11, bold=True)
            vr = p.add_run(", ".join(items) if isinstance(items, list) else str(items))
            _set_run(vr, 11)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)

    # Certifications
    certs = d.get("certifications", [])
    if certs:
        _section_heading("Certifications")
        for cert in certs:
            parts = [cert.get("name",""), cert.get("issuer","")]
            if cert.get("year"): parts.append(str(cert["year"]))
            _bullet("  |  ".join(p for p in parts if p))

    buf = BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()
    logger.info("DOCX generated: %d bytes", len(docx_bytes))
    return docx_bytes


# ─────────────────────────────────────────────────────────────────────────────
# PDF EXPORT — reportlab, fully ATS-safe (NO Table layout)
# ─────────────────────────────────────────────────────────────────────────────
# ATS rule: Never use Table for layout. All text must flow left-to-right in
# a single column so PDF text-extraction reads in natural reading order.
# Right-aligned dates are achieved with a right tab stop on a single paragraph.
# ─────────────────────────────────────────────────────────────────────────────

def _export_pdf(d: dict) -> bytes:
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, KeepTogether
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
        from reportlab.lib.units import inch
        from reportlab.platypus.flowables import Flowable
    except ImportError:
        raise RuntimeError("reportlab not installed. Run: pip install reportlab")

    BLACK = colors.HexColor("#000000")
    GREY  = colors.HexColor("#444444")

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=36, rightMargin=36,
        topMargin=36,  bottomMargin=36,
    )

    # Page content width = 8.5in - 0.5in - 0.5in = 7.5in = 540pt
    PAGE_W = 540

    # ── Styles ────────────────────────────────────────────────────────────────
    NAME    = ParagraphStyle("Name",    fontName="Helvetica-Bold",    fontSize=22, alignment=TA_CENTER, spaceAfter=2,  textColor=BLACK)
    CONTACT = ParagraphStyle("Contact", fontName="Helvetica",         fontSize=10, alignment=TA_CENTER, spaceAfter=10, textColor=BLACK)
    SEC_H   = ParagraphStyle("SecH",    fontName="Helvetica-Bold",    fontSize=12, spaceBefore=8, spaceAfter=1, textColor=BLACK)
    BODY    = ParagraphStyle("Body",    fontName="Helvetica",         fontSize=11, spaceBefore=2, spaceAfter=2, leading=14, alignment=TA_JUSTIFY, textColor=BLACK)
    ROLE_L  = ParagraphStyle("AreaL",   fontName="Helvetica-Bold",    fontSize=11, spaceBefore=5, spaceAfter=1, leading=14, textColor=BLACK)
    # Bullet: hanging indent so wrapped lines align under the text, not the bullet
    BULLET  = ParagraphStyle("Bullet",  fontName="Helvetica",         fontSize=11, spaceBefore=1, spaceAfter=1, leading=14,
                              leftIndent=14, firstLineIndent=-14, textColor=BLACK)
    SCHOOL  = ParagraphStyle("School",  fontName="Helvetica-Oblique", fontSize=10, spaceBefore=2, spaceAfter=4, textColor=GREY)

    def safe(t: str) -> str:
        return str(t).replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")

    def div() -> HRFlowable:
        return HRFlowable(width="100%", thickness=0.5, color=BLACK, spaceBefore=1, spaceAfter=4)

    def sh(title: str) -> list:
        """Section heading + horizontal rule."""
        return [Paragraph(safe(title.upper()), SEC_H), div()]

    def bul(text: str) -> Paragraph:
        """ATS-safe bullet: standard white circle (Jerom Tom style)"""
        return Paragraph(f"\u25CB  {safe(text)}", BULLET)

    # ── Right-tab paragraph for "Role / Title    DATE" on one line ────────────
    # ATS reads this as: "Software Engineer — Acme Corp    2023-Present"
    # which is correct natural order. No table involved.
    from reportlab.lib.styles import ParagraphStyle as _PS
    from reportlab.platypus import Paragraph as _P

    def role_line(left_bold: str, right_italic: str) -> Paragraph:
        """
        Single paragraph: bold left text + right-aligned date via XML tab stop.
        ATS text extraction reads: "left_bold    right_italic" — correct order.
        """
        left_clean  = safe(left_bold)
        right_clean = safe(right_italic)
        if right_clean:
            text = f"<b>{left_clean}</b>    <i>{right_clean}</i>"
        else:
            text = f"<b>{left_clean}</b>"
        return Paragraph(text, ROLE_L)

    story = []

    # ── Name + Contact ────────────────────────────────────────────────────────
    story.append(Paragraph(safe(d.get("name", "Candidate")), NAME))
    if d.get("contact"):
        story.append(Paragraph(safe(d["contact"]).replace("  |  ", "  |  "), CONTACT))

    # ── Summary ───────────────────────────────────────────────────────────────
    obj = (d.get("objective") or "").strip()
    if obj:
        story += sh("Summary")
        story.append(Paragraph(safe(obj), BODY))
        story.append(Spacer(1, 4))

    # ── Work Experience ───────────────────────────────────────────────────────
    exp_list = d.get("experience", [])
    if exp_list:
        story += sh("Experience")
        for exp in exp_list:
            role = exp.get("role", "")
            co   = exp.get("company", "")
            dur  = exp.get("duration", "")
            left = role + (f" — {co}" if co else "")
            block = [role_line(left, dur)]
            for r in exp.get("responsibilities", []):
                if r and r.strip():
                    block.append(bul(r.strip()))
            block.append(Spacer(1, 4))
            story.append(KeepTogether(block))

    # ── Projects ──────────────────────────────────────────────────────────────
    proj_list = d.get("projects", [])
    if proj_list:
        story += sh("Projects")
        for proj in proj_list:
            title = proj.get("title", "")
            techs = proj.get("technologies", [])
            right = ", ".join(techs) if techs else ""
            block = [role_line(title, right)]
            desc  = (proj.get("description") or "").strip()
            if desc:
                block.append(bul(desc))
            for m in proj.get("metrics", []):
                if m and m.strip():
                    block.append(bul(m.strip()))
            block.append(Spacer(1, 4))
            story.append(KeepTogether(block))

    # ── Education ─────────────────────────────────────────────────────────────
    education = d.get("education", [])
    if education:
        story += sh("Education")
        import re
        for edu in education:
            inst  = edu.get("institution", "")
            deg   = edu.get("degree", "")
            gpa   = edu.get("gpa", "")
            yr    = str(edu.get("graduation_year", ""))
            level = str(edu.get("level", "")).lower()

            is_school = (level in ("class12","class10") or
                         bool(re.search(r"class\s*(x|xii?|10|12)", inst + " " + deg, re.I)))

            if is_school:
                # Line 1: School Name + Year
                story.append(role_line(inst or deg, yr))
                # Line 2: Grade label + percentage (lighter gray)
                grade_label = "Class XII" if "12" in deg + inst else "Class X"
                if gpa:
                    story.append(Paragraph(safe(f"{grade_label}  |  {gpa}"), SCHOOL))
            else:
                # Line 1: Institution Name + Year
                story.append(role_line(inst or deg, yr))
                # Line 2: Course + CGPA
                left2_parts = [p for p in [deg, f"CGPA: {gpa}" if gpa else ""] if p]
                left2 = "  |  ".join(left2_parts)
                if left2:
                    story.append(Paragraph(safe(left2), SCHOOL))
            story.append(Spacer(1, 6))

        # School 10th / 12th (from non-LLM fallback path)
        s10 = d.get("school_10th")
        s12 = d.get("school_12th")
        if s10 or s12:
            if s12:
                story.append(role_line("Class XII", ""))
                story.append(Paragraph(safe(f"{s12}"), SCHOOL))
            if s10:
                story.append(role_line("Class X", ""))
                story.append(Paragraph(safe(f"{s10}"), SCHOOL))
            story.append(Spacer(1, 4))

    # ── Technical Skills ──────────────────────────────────────────────────────
    skills = {k: v for k, v in d.get("skills", {}).items() if v}
    if skills:
        story += sh("Skills")
        for cat, items in skills.items():
            iv = ", ".join(items) if isinstance(items, list) else str(items)
            if iv.strip():
                story.append(Paragraph(f"<b>{safe(cat)}:</b> {safe(iv)}", BODY))
        story.append(Spacer(1, 4))

    # ── Certifications ────────────────────────────────────────────────────────
    certs = d.get("certifications", [])
    if certs:
        story += sh("Certifications")
        for cert in certs:
            parts = [p for p in [cert.get("name",""), cert.get("issuer",""), str(cert.get("year","") or "")] if p]
            story.append(bul("  |  ".join(parts)))

    doc.build(story)
    pdf_bytes = buf.getvalue()
    logger.info("PDF generated (ATS-safe, no Table): %d bytes", len(pdf_bytes))
    return pdf_bytes


# ─────────────────────────────────────────────────────────────────────────────
# LATEX → PDF via latexonline.cc
# ─────────────────────────────────────────────────────────────────────────────

def compile_latex_to_pdf(latex_code: str) -> bytes:
    import requests
    
    # Official TeX Users Group public compiler
    url = "https://texlive.net/cgi-bin/latexcgi"
    logger.info("Compiling LaTeX via texlive.net (%d chars)", len(latex_code))

    try:
        files = {
            'filecontents[]': ('document.tex', latex_code, 'application/x-tex')
        }
        data = {
            'filename[]': 'document.tex',
            'engine': 'pdflatex',
            'return': 'pdf'
        }
        
        response = requests.post(url, files=files, data=data, timeout=60)
        
        if response.status_code != 200:
            logger.error("LaTeX API Error %d: %s", response.status_code, response.text[:200])
            raise RuntimeError(f"LaTeX API returned code {response.status_code}")
            
        pdf_bytes = response.content
        if not pdf_bytes.startswith(b"%PDF"):
            # texlive.net returns the compilation logs if it fails!
            sample = pdf_bytes.decode(errors='ignore')
            logger.error("LaTeX Compilation Log/Error:\n%s", sample[-4000:])
            raise RuntimeError("LaTeX document contains syntax errors — compilation failed")
            
        logger.info("LaTeX compilation successful: %d bytes", len(pdf_bytes))
        return pdf_bytes

    except Exception as e:
        logger.error("LaTeX compilation failed: %s", e)
        raise RuntimeError(f"LaTeX compilation failed: {e}")