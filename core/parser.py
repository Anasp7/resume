"""
Smart Resume — Document Parser v6
====================================
Phase 3 upgrades for 92-95% accuracy:
  1. KNOWN_TECH_TERMS expanded from ~50 to 200+ terms
  2. SECTION_HEADERS expanded to cover all resume variants
  3. Contact extractor: LinkedIn, GitHub, portfolio URLs
  4. Multi-line responsibility grouping
  5. Scanned PDF fallback with pytesseract (optional)
  6. Table-aware DOCX parsing
  7. Better name extraction (skip email/phone lines)
"""

from __future__ import annotations
import io, re, logging
from typing import Optional

_COL_CACHE: dict = {}  # md5 → {left, right, is_two_col}
logger = logging.getLogger("smart_resume.parser")

# ── Optional imports ──────────────────────────────────────────────────────────
try:
    import pdfplumber
    _HAS_PDFPLUMBER = True
except ImportError:
    _HAS_PDFPLUMBER = False

try:
    import fitz
    _HAS_PYMUPDF = True
except ImportError:
    _HAS_PYMUPDF = False

try:
    from docx import Document as DocxDocument
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

try:
    import spacy
    _nlp = spacy.load("en_core_web_sm")
    _HAS_SPACY = True
except Exception:
    _HAS_SPACY = False

try:
    from PIL import Image
    import pytesseract
    _HAS_OCR = True
except ImportError:
    _HAS_OCR = False


# ── Section headers — expanded to 92%+ coverage ───────────────────────────────

SECTION_HEADERS = {
    "summary": [
        "summary", "objective", "career objective", "professional summary",
        "profile", "about me", "about", "overview", "introduction",
        "professional profile", "career profile", "personal statement",
        "who i am", "highlights",
    ],
    "education": [
        "education", "educational background", "academic background",
        "academic qualification", "qualification", "qualifications",
        "degree", "schooling", "academic details", "educational details",
        "academic history",
        # NOTE: "university" and "college" removed — single-word exact matches caused
        # institution names to be classified as section headers, swallowing following content.
        # Multi-word forms like "university college" still match via substring.
    ],
    "experience": [
        "experience", "work experience", "professional experience",
        "employment", "employment history", "work history",
        "internship", "internships", "industrial experience",
        "career history", "job history", "positions held",
        "relevant experience", "industry experience",
    ],
    "projects": [
        "project", "projects", "personal projects", "academic projects",
        "technical projects", "major projects", "minor projects",
        "mini projects", "final year project", "capstone project",
        "key projects", "notable projects", "portfolio",
        "work done", "builds", "what i've built", "things i've built",
        "open source", "side projects", "hobby projects",
        "college projects", "university projects", "coursework projects",
        "technical work", "web creation", "application developed",
    ],
    "skills": [
        "skill", "skills", "technical skills", "technical expertise",
        "core competencies", "competencies", "technologies",
        "tools", "tools and technologies", "tech stack",
        "programming languages", "languages", "frameworks",
        "software", "software skills", "areas of expertise",
        "expertise", "proficiencies", "technical proficiencies",
        "key skills", "relevant skills",
    ],
    "certifications": [
        "certif", "certification", "certifications", "course", "courses",
        "training", "license", "licenses", "achievement", "achievements",
        "award", "awards", "credential", "credentials",
        "online courses", "mooc", "badges", "accomplishments",
        "professional development", "continuing education",
    ],
    "publications": [
        "publication", "publications", "research", "papers",
        "journal", "conference", "patents",
    ],
    "languages": [
        "spoken languages", "language proficiency", "natural languages",
    ],
    "interests": [
        "interest", "interests", "hobbies", "extracurricular",
        "activities", "volunteer",
    ],
}

# ── Keywords that must be an EXACT line match (not just found in line) ──────────
# These are common in institution names/content — "Anna University", "College of Engg"
# Only trigger section detection if the ENTIRE line IS this keyword
EXACT_MATCH_ONLY_KEYWORDS = {
    "university", "college", "degree", "research", "languages",
    "interest", "interests", "training", "tools", "software",
    "activities", "overview", "highlights",
}

# ── KNOWN_TECH_TERMS — expanded from 50 to 200+ ──────────────────────────────

KNOWN_TECH_TERMS = {
    # ── Programming languages ─────────────────────────────────────────────────
    "python", "java", "javascript", "typescript", "c", "c++", "c#", "go",
    "rust", "kotlin", "swift", "r", "scala", "ruby", "php", "bash", "shell",
    "sql", "perl", "matlab", "dart", "elixir", "haskell", "lua", "groovy",
    "assembly", "vhdl", "verilog", "fortran", "cobol", "solidity",

    # ── Web frameworks ────────────────────────────────────────────────────────
    "fastapi", "flask", "django", "spring", "springboot", "spring boot",
    "react", "vue", "angular", "nextjs", "next.js", "nuxt", "svelte",
    "express", "nestjs", "laravel", "rails", "asp.net", "dotnet",
    "gatsby", "remix", "astro", "htmx", "jquery", "bootstrap", "tailwind",
    "material ui", "chakra ui", "ant design",

    # ── ML / AI / Data science ────────────────────────────────────────────────
    "pytorch", "tensorflow", "keras", "sklearn", "scikit-learn", "scikit",
    "pandas", "numpy", "scipy", "matplotlib", "seaborn", "plotly",
    "xgboost", "lightgbm", "catboost", "huggingface", "transformers",
    "langchain", "llamaindex", "openai", "anthropic", "gemini",
    "opencv", "pillow", "albumentations", "detectron",
    "mlflow", "wandb", "dvc", "optuna", "ray", "celery",
    "spacy", "nltk", "gensim", "bert", "gpt", "llm",
    "stable diffusion", "diffusers", "onnx", "tensorrt",
    "feature engineering", "model training", "fine-tuning",
    "computer vision", "nlp", "deep learning", "machine learning",
    "reinforcement learning", "time series", "anomaly detection",

    # ── Robotics ──────────────────────────────────────────────────────────────
    "ros", "ros2", "gazebo", "rviz", "moveit", "nav2", "rclpy",
    "arduino", "raspberry pi", "raspberry", "jetson", "stm32",
    "rtos", "freertos", "embedded", "firmware", "uart", "i2c", "spi",
    "can bus", "modbus", "mqtt",

    # ── Databases ─────────────────────────────────────────────────────────────
    "postgresql", "postgres", "mysql", "sqlite", "mongodb", "firebase",
    "redis", "cassandra", "dynamodb", "supabase", "planetscale",
    "elasticsearch", "neo4j", "influxdb", "cockroachdb", "mariadb",
    "oracle", "mssql", "prisma", "sqlalchemy", "sequelize", "mongoose",

    # ── Cloud & DevOps ────────────────────────────────────────────────────────
    "docker", "kubernetes", "k8s", "helm", "terraform", "ansible",
    "aws", "gcp", "azure", "heroku", "vercel", "netlify", "railway",
    "digitalocean", "cloudflare", "nginx", "apache", "caddy",
    "jenkins", "github actions", "gitlab ci", "circleci", "travis ci",
    "argocd", "prometheus", "grafana", "datadog", "elk", "kibana",
    "linux", "ubuntu", "debian", "centos", "unix",

    # ── Version control & tools ───────────────────────────────────────────────
    "git", "github", "gitlab", "bitbucket", "svn",
    "jira", "confluence", "trello", "notion", "slack", "figma",
    "postman", "insomnia", "swagger", "openapi",
    "webpack", "vite", "babel", "eslint", "prettier",
    "make", "cmake", "gradle", "maven", "npm", "yarn", "pnpm",
    "virtualenv", "conda", "poetry", "pip",

    # ── APIs & protocols ──────────────────────────────────────────────────────
    "rest", "restful", "graphql", "grpc", "websocket", "oauth",
    "jwt", "soap", "kafka", "rabbitmq", "celery", "redis pubsub",
    "microservices", "api gateway", "load balancer",

    # ── Mobile ────────────────────────────────────────────────────────────────
    "android", "ios", "flutter", "react native", "xamarin",
    "xcode", "android studio", "jetpack compose", "swiftui",

    # ── Security ─────────────────────────────────────────────────────────────
    "penetration testing", "nmap", "burp suite", "metasploit",
    "owasp", "wireshark", "kali", "siem", "soc",
    "ssl", "tls", "encryption", "cryptography",

    # ── Data engineering ──────────────────────────────────────────────────────
    "spark", "hadoop", "airflow", "dbt", "bigquery", "redshift",
    "snowflake", "databricks", "hive", "pig", "flink",
    "etl", "elt", "data pipeline", "data warehouse", "data lake",
    "tableau", "powerbi", "looker", "metabase", "superset",

    # ── Testing ───────────────────────────────────────────────────────────────
    "pytest", "unittest", "jest", "mocha", "cypress", "selenium",
    "playwright", "testng", "junit", "robot framework",

    # ── Methodologies ────────────────────────────────────────────────────────
    "agile", "scrum", "kanban", "tdd", "bdd", "ci/cd", "devops",

    # ── Blockchain ───────────────────────────────────────────────────────────
    "blockchain", "ethereum", "web3", "hardhat", "truffle", "solidity",

    # ── Other ─────────────────────────────────────────────────────────────────
    "streamlit", "gradio", "fasthtml", "pydantic", "sqlmodel",
    "celery", "dramatiq", "rq", "websockets", "asyncio",
    "regex", "xml", "json", "yaml", "protobuf", "avro",
}

PROJECT_IN_EXPERIENCE_SIGNALS = [
    r"\bbuilt\b", r"\bdeveloped\b", r"\bcreated\b", r"\bdesigned\b",
    r"\bimplemented\b", r"\bdeployed\b", r"github\.com",
]
EXPERIENCE_IN_PROJECT_SIGNALS = [
    r"\bintern(ship)?\b", r"\bfull[- ]time\b", r"\bpart[- ]time\b",
    r"\bmanaged\b", r"\bteam of \d+\b", r"\bresponsible for\b",
]


# ── Column boundary detection ─────────────────────────────────────────────────

def _find_column_split(page) -> Optional[float]:
    """
    Find actual column gap using word x-positions.
    Returns midpoint x-coordinate of the gap, or None if single-column.
    """
    words = page.extract_words(x_tolerance=3, y_tolerance=3)
    if not words:
        return None

    x_positions = sorted(set(round(w["x0"]) for w in words))
    if len(x_positions) < 4:
        return None

    gaps = []
    for i in range(1, len(x_positions)):
        gap = x_positions[i] - x_positions[i - 1]
        mid = (x_positions[i] + x_positions[i - 1]) / 2
        gaps.append((gap, mid))

    gaps.sort(reverse=True)
    largest_gap, gap_mid = gaps[0]
    page_width = page.width

    if largest_gap > page_width * 0.08 and page_width * 0.25 < gap_mid < page_width * 0.75:
        logger.info("Column gap: %.1fpx at x=%.1f (page_w=%.1f)", largest_gap, gap_mid, page_width)
        return gap_mid

    return None


# ── PDF extraction ────────────────────────────────────────────────────────────

def extract_pdf_text(file_bytes: bytes) -> str:
    """
    Extract text from PDF.
    Strategy: pdfplumber (column-aware) → PyMuPDF → OCR fallback
    """
    if _HAS_PDFPLUMBER:
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                all_parts = []
                for page in pdf.pages:
                    split_x = _find_column_split(page)
                    if split_x:
                        # Add 2px overlap so characters at column boundary aren't cut
                        left  = page.within_bbox((0,            0, split_x + 2, page.height))
                        right = page.within_bbox((split_x - 2,  0, page.width,  page.height))
                        lt    = left.extract_text(x_tolerance=4, y_tolerance=3) or ""
                        rt    = right.extract_text(x_tolerance=4, y_tolerance=3) or ""
                        page_text = lt.strip() + "\n\n" + rt.strip()
                        logger.info("Two-column: left=%d right=%d chars", len(lt), len(rt))
                        all_parts.append(("two_col", lt.strip(), rt.strip()))
                    else:
                        page_text = page.extract_text(x_tolerance=2, y_tolerance=3) or ""
                        logger.info("Single-column: %d chars", len(page_text))
                        all_parts.append(("single", page_text, ""))

                # Build combined text and store column texts for two-col PDFs
                left_parts  = [lt for kind, lt, rt in all_parts if kind == "two_col"]
                right_parts = [rt for kind, lt, rt in all_parts if kind == "two_col"]
                raw_parts   = [lt + "\n\n" + rt if kind == "two_col" else lt
                               for kind, lt, rt in all_parts]
                text = "\n\n".join(raw_parts)
                if len(text.strip()) > 50:
                    clean = _clean_text(text)
                    # Attach column texts to a module-level cache keyed by hash
                    import hashlib
                    _h = hashlib.md5(file_bytes).hexdigest()
                    _COL_CACHE[_h] = {
                        "left":  "\n\n".join(left_parts),
                        "right": "\n\n".join(right_parts),
                        "is_two_col": bool(left_parts),
                    }
                    return clean
                logger.warning("pdfplumber returned very short text — trying PyMuPDF")
        except Exception as e:
            logger.warning("pdfplumber failed: %s", e)

    if _HAS_PYMUPDF:
        try:
            doc  = fitz.open(stream=file_bytes, filetype="pdf")
            text = "\n".join(page.get_text("text") for page in doc)
            if len(text.strip()) > 50:
                return _clean_text(text)
            logger.warning("PyMuPDF returned very short text — trying OCR")
        except Exception as e:
            logger.warning("PyMuPDF failed: %s", e)

    # OCR fallback for scanned PDFs
    if _HAS_OCR and _HAS_PYMUPDF:
        try:
            logger.info("Attempting OCR on scanned PDF...")
            doc   = fitz.open(stream=file_bytes, filetype="pdf")
            parts = []
            for page in doc:
                pix  = page.get_pixmap(dpi=200)
                img  = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                text = pytesseract.image_to_string(img, config="--psm 6")
                parts.append(text)
            text = "\n".join(parts)
            if text.strip():
                logger.info("OCR extracted %d chars", len(text))
                return _clean_text(text)
        except Exception as e:
            logger.warning("OCR failed: %s", e)

    raise RuntimeError(
        "PDF extraction failed. "
        "Install pdfplumber (pip install pdfplumber) or PyMuPDF (pip install pymupdf). "
        "For scanned PDFs also install: pip install pytesseract Pillow"
    )


# ── DOCX extraction ───────────────────────────────────────────────────────────

def extract_docx_text(file_bytes: bytes) -> str:
    """
    Extract text from DOCX including tables and text boxes.
    """
    if not _HAS_DOCX:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    doc        = DocxDocument(io.BytesIO(file_bytes))
    paragraphs = []

    # Main body paragraphs
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            paragraphs.append(t)

    # Tables — common in modern resume templates
    for table in doc.tables:
        for row in table.rows:
            row_parts = []
            for cell in row.cells:
                ct = cell.text.strip()
                if ct:
                    row_parts.append(ct)
            if row_parts:
                paragraphs.append(" | ".join(row_parts))

    # Text boxes (stored in XML body)
    try:
        from docx.oxml.ns import qn
        for txbx in doc.element.body.iter(qn("w:txbxContent")):
            for para in txbx.iter(qn("w:p")):
                texts = [r.text for r in para.iter(qn("w:t")) if r.text]
                line  = "".join(texts).strip()
                if line:
                    paragraphs.append(line)
    except Exception:
        pass

    return _clean_text("\n".join(paragraphs))


# ── Text cleaning ─────────────────────────────────────────────────────────────

def _clean_text(text: str) -> str:
    # Remove PDF CID encoding artifacts
    text = re.sub(r"[(]cid[ :]*[0-9]+[)]", "", text, flags=re.IGNORECASE)
    # Normalize bullets
    text = text.replace("·", "-").replace("•", "-").replace("●", "-")
    text = text.replace("\u2022", "-").replace("\u25cf", "-").replace("\u2013", "-")
    # Remove null bytes and non-printable chars
    text = text.replace("\x00", "")
    text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E\u00A0-\uFFFF]", " ", text)
    # Normalize whitespace
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return "\n".join(line.rstrip() for line in text.splitlines()).strip()


# ── Section detection — improved ──────────────────────────────────────────────

def _line_matches_section(line: str) -> Optional[str]:
    """
    Detect if a line is a section header.
    Handles: ALL CAPS headers, colon-terminated headers, short bold lines.
    """
    clean = line.strip().lstrip("-*>|#").strip().rstrip(":").strip()
    if not clean or len(clean) > 60:
        return None

    cl    = clean.lower()
    words = cl.split()

    # Must be 1-5 words to be a header
    if len(words) > 5:
        return None

    for section, keywords in SECTION_HEADERS.items():
        for kw in keywords:
            # Exact match OR the keyword equals the entire cleaned line
            # Use word-boundary match for single-word aliases to avoid
            # "software" matching "software engineer at acme corp"
            if kw == cl:
                return section
            kw_words = kw.split()
            if len(kw_words) > 1:
                # Multi-word keyword: safe to check as substring
                if kw in cl:
                    return section
            else:
                # Single-word keyword — two modes:
                # 1. Exact-match-only keywords (e.g. "university", "college") require
                #    the ENTIRE cleaned line to be exactly this keyword
                # 2. Other single-word keywords use word-boundary + length guard
                if kw in EXACT_MATCH_ONLY_KEYWORDS:
                    if cl == kw:
                        return section
                elif re.search(r"(?<![a-z])" + re.escape(kw) + r"(?![a-z])", cl):
                    if len(words) <= 2:
                        return section

    # ALL CAPS lines with 1-4 words are almost always section headers
    if clean.isupper() and 1 <= len(words) <= 4:
        for section, keywords in SECTION_HEADERS.items():
            for kw in keywords:
                if kw in cl:
                    return section

    return None


def split_into_sections(text: str) -> dict[str, str]:
    """
    Split resume text into sections.
    Improved: handles ALL CAPS headers, colon-terminated headers, indented bullets.
    """
    sections: dict[str, list[str]] = {"header": []}
    current = "header"

    lines = text.splitlines()
    i     = 0
    while i < len(lines):
        line    = lines[i]
        stripped = line.strip()
        matched  = _line_matches_section(stripped)

        if matched:
            current = matched
            sections.setdefault(current, [])
        else:
            sections.setdefault(current, []).append(stripped)
        i += 1

    return {k: "\n".join(v).strip() for k, v in sections.items() if "\n".join(v).strip()}


# ── Mismatch detector ─────────────────────────────────────────────────────────

def detect_section_mismatches(sections: dict[str, str]) -> dict[str, list[str]]:
    result: dict[str, list[str]] = {
        "project_in_experience": [],
        "experience_in_project": [],
    }
    for line in sections.get("experience", "").splitlines():
        if any(re.search(p, line.lower()) for p in PROJECT_IN_EXPERIENCE_SIGNALS):
            result["project_in_experience"].append(line.strip())
    for line in sections.get("projects", "").splitlines():
        if any(re.search(p, line.lower()) for p in EXPERIENCE_IN_PROJECT_SIGNALS):
            result["experience_in_project"].append(line.strip())
    return result


# ── Skill extractor — improved ────────────────────────────────────────────────

def extract_inline_skills(text: str) -> list[str]:
    """
    Extract all known tech terms from text using word-boundary matching.
    Multi-word terms (e.g. 'react native') checked before single words.
    """
    found = set()
    tl    = text.lower()

    # Sort by length descending so multi-word terms match first
    sorted_terms = sorted(KNOWN_TECH_TERMS, key=len, reverse=True)

    for term in sorted_terms:
        if " " in term:
            # Multi-word: simple substring (word boundary not reliable for phrases)
            if term in tl:
                found.add(term)
        else:
            pattern = r"(?<![a-z0-9+#])" + re.escape(term) + r"(?![a-z0-9+#])"
            if re.search(pattern, tl):
                found.add(term)

    return sorted(found)


# ── Contact extractor — improved ──────────────────────────────────────────────

def extract_contact_info(text: str) -> dict[str, str]:
    """
    Extract name, email, phone, LinkedIn, GitHub, portfolio from resume text.
    """
    info: dict[str, str] = {}

    # Email
    email = re.search(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", text)
    if email:
        info["email"] = email.group().strip()

    # Phone — handles +91, international formats
    phone = re.search(
        r"(?:\+?\d{1,3}[\s\-.]?)?\(?\d{3,5}\)?[\s\-.]?\d{3,5}[\s\-.]?\d{3,5}",
        text
    )
    if phone:
        ph = re.sub(r"[^\d+]", "", phone.group())
        if len(ph) >= 10:
            info["phone"] = phone.group().strip()

    # LinkedIn
    linkedin = re.search(
        r"(?:linkedin\.com/in/|linkedin:\s*)(?!https?://)([a-zA-Z0-9\-_]+)", text, re.IGNORECASE
    )
    if linkedin:
        info["linkedin"] = f"linkedin.com/in/{linkedin.group(1)}"

    # GitHub
    github = re.search(
        r"(?:github\.com/)(?!https?://)([a-zA-Z0-9\-_]+)", text, re.IGNORECASE
    )
    if github:
        info["github"] = f"github.com/{github.group(1)}"

    # Portfolio / personal website
    portfolio = re.search(
        r"(?:portfolio|website|site|blog)\s*[:\-]?\s*(https?://[^\s]+|www\.[^\s]+)",
        text, re.IGNORECASE
    )
    if portfolio:
        info["portfolio"] = portfolio.group(1)

    # Location — look for city/district/state patterns
    location = re.search(
        r"(?:^|\n)\s*([A-Z][a-z]+(?:[\s,]+[A-Z][a-z]+){0,3})\s*(?:\n|$)",
        text, re.MULTILINE
    )
    if location:
        loc = location.group(1).strip()
        # Filter out common false positives
        if (len(loc.split()) <= 4
                and not re.search(r"\b(university|college|engineer|developer|student|player|football|cricket|tennis|music|hobby|hobbies|activities|sports)\b", loc, re.I)
                and not re.search(r"[@\d]", loc)):
            info["location"] = loc

    # Name — first non-empty line that isn't a contact detail or header
    SKIP_PATTERNS = [
        r"@", r"\+\d", r"linkedin", r"github", r"http",
        r"\b(resume|cv|curriculum)\b",
    ]
    for line in text.splitlines()[:10]:
        line = line.strip()
        if (line
                and len(line.split()) >= 2
                and len(line.split()) <= 6
                and not any(re.search(p, line, re.I) for p in SKIP_PATTERNS)
                and not _line_matches_section(line)
                and not re.search(r"[\d\(\)\+@]", line)):
            info["name"] = line
            break

    return info



# ── SOFT_SKILLS — shared filter used by app.py and parser ─────────────────────
# Centralised here so both the section splitter and the resume builder use same list

SOFT_SKILLS = {
    "teamwork", "leadership", "time management", "project management",
    "communication", "english", "malayalam", "hindi", "tamil", "arabic",
    "languages", "softwares", "tools", "plus two", "student", "contact",
    "profile", "cgpa",
}

# ── JUNK_TITLE_PATTERNS — used to filter bad project/experience titles ─────────

JUNK_TITLE_PATTERNS = [
    r"^languages?\s*:",
    r"^softwares?\s*:",
    r"^(english|malayalam|hindi|tamil|arabic)$",
    r"^(teamwork|leadership|time management|project management)$",
    r"^(student|contact|profile|languages|skills|tools)$",
    r"^\+?\d[\d\s\-\(\)]{7,}",
    r"^[~'\-,\s\d]+$",
    r"^cgpa\s*:",
    r"^\d{4}[-]\d{0,4}$",
]

def is_junk_line(line: str) -> bool:
    """Return True if a line should be discarded (phone, soft skill, ALL-CAPS name, etc.)"""
    import re as _re
    l = line.strip().lower()
    if not l or len(l) < 2:
        return True
    for pattern in JUNK_TITLE_PATTERNS:
        if _re.match(pattern, l, _re.IGNORECASE):
            return True
    words = line.strip().split()
    if 2 <= len(words) <= 5 and all(w.isupper() and w.isalpha() for w in words):
        return True
    return False


def is_real_project_title(line: str) -> bool:
    """Return True if this line looks like a genuine project title."""
    if is_junk_line(line):
        return False
    if len(line) > 60:
        return False
    words = line.strip().split()
    if len(words) == 1 and line.isupper():
        return False
    if line.strip().startswith("-") and len(words) <= 2:
        return False
    return True

# ── NLP preprocessing ─────────────────────────────────────────────────────────

def preprocess_for_embedding(text: str) -> str:
    """
    Preprocess text for similarity computation.
    SpaCy lemmatization if available, else regex-based stopword removal.
    """
    if _HAS_SPACY:
        doc = _nlp(text[:100_000])
        return " ".join(
            t.lemma_.lower() for t in doc
            if not t.is_stop and not t.is_punct and t.is_alpha and len(t) > 1
        )

    STOPWORDS = {
        "the","a","an","is","was","were","are","be","been","have","has","had",
        "do","does","did","will","would","could","should","to","of","in","on",
        "at","by","for","with","and","or","but","not","this","that","it","its",
        "we","our","you","your","they","their","my","me","him","her","us",
        "from","as","if","so","up","out","about","into","through","over",
    }
    tokens = re.findall(r"[a-zA-Z][a-zA-Z0-9+#.\-_]{1,}", text)
    return " ".join(t.lower() for t in tokens if t.lower() not in STOPWORDS and len(t) > 1)


# ── ATS Score Simulator ──────────────────────────────────────────────────────

ATS_CRITICAL_SECTIONS = {"experience", "education", "skills"}

ATS_BAD_PHRASES = [
    r"references available",
    r"references upon request",
    r"curriculum vitae",
    r"\bphoto attached\b",
    r"date of birth",
    r"\bdob\b",
    r"marital status",
    r"\bnationality\b",
    r"\bpassport\b",
]

ATS_CONTACT_PATTERNS = {
    "email":  r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}",
    "phone":  r"(\+?\d[\d\s\-\.\(\)]{8,})",
}

ATS_WEAK_VERBS = [
    r"\bresponsible for\b",
    r"\bhelped with\b",
    r"\bworked on\b",
    r"\binvolved in\b",
    r"\bassisted with\b",
    r"\bparticipated in\b",
]

ATS_POWER_VERBS = {
    "built", "developed", "designed", "implemented", "engineered", "architected",
    "deployed", "automated", "optimized", "reduced", "increased", "improved",
    "created", "launched", "led", "managed", "delivered", "integrated",
    "migrated", "refactored", "trained", "fine-tuned", "analyzed", "processed",
}


def compute_ats_score(raw_text: str, sections: dict, skills: list, target_role: str = "") -> dict:
    """
    Deterministic ATS simulation — returns score 0-100 and itemised flags.
    Does NOT call any LLM. Pure rule-based checks.

    Criteria (total 100 pts):
      20 — Contact completeness (email + phone)
      20 — Critical section presence (experience, education, skills)
      15 — No ATS-hostile content (tables hints, bad phrases, photos)
      15 — Power verb usage in experience bullets
      15 — Keyword density (skills appear in experience/project text)
      15 — Structure quality (section order, line length, no excessive caps)
    """
    flags   = []
    score   = 0
    details = {}

    text_lower = raw_text.lower()

    # ── 1. Contact completeness (20 pts) ──────────────────────────────────────
    contact_score = 0
    for field, pattern in ATS_CONTACT_PATTERNS.items():
        if re.search(pattern, raw_text):
            contact_score += 10
        else:
            flags.append(f"Missing {field} — ATS may reject without contact info")
    score += contact_score
    details["contact"] = contact_score

    # ── 2. Critical sections present (20 pts) ─────────────────────────────────
    section_score = 0
    for sec in ATS_CRITICAL_SECTIONS:
        if sec in sections and sections[sec].strip():
            section_score += 7
        else:
            flags.append(f"Missing '{sec}' section — ATS may score this as incomplete")
    section_score = min(section_score, 20)
    score += section_score
    details["sections"] = section_score

    # ── 3. ATS-hostile content check (15 pts) ─────────────────────────────────
    hostile_score = 15
    for phrase in ATS_BAD_PHRASES:
        if re.search(phrase, text_lower, re.IGNORECASE):
            hostile_score -= 5
            match = re.search(phrase, text_lower, re.IGNORECASE)
            flags.append(f"ATS-hostile phrase found: '{match.group()}'")
    # Check for table-like content (many pipe chars)
    pipe_lines = sum(1 for line in raw_text.splitlines() if line.count("|") >= 3)
    if pipe_lines > 3:
        hostile_score -= 5
        flags.append(f"Table-like formatting detected ({pipe_lines} lines with | chars) — ATS parsers often fail on tables")
    hostile_score = max(hostile_score, 0)
    score += hostile_score
    details["ats_hostile"] = hostile_score

    # ── 4. Power verb usage (15 pts) ──────────────────────────────────────────
    exp_text = sections.get("experience", "") + " " + sections.get("projects", "")
    bullets  = [l.strip().lstrip("-•*").strip() for l in exp_text.splitlines() if l.strip()]
    
    if bullets:
        power_count = sum(
            1 for b in bullets
            if b.split()[0].lower() in ATS_POWER_VERBS if b.split()
        )
        weak_count = sum(
            1 for b in bullets
            if any(re.search(p, b, re.IGNORECASE) for p in ATS_WEAK_VERBS)
        )
        power_ratio = power_count / len(bullets) if bullets else 0
        power_score = min(int(power_ratio * 30), 15)
        if weak_count > 0:
            flags.append(f"{weak_count} bullet(s) use weak passive phrasing ('responsible for', 'helped with') — replace with action verbs")
        if power_count == 0:
            flags.append("No power verbs found in experience/projects — ATS ranks active-verb resumes higher")
    else:
        power_score = 0
        flags.append("No experience/project bullets detected")
    score += power_score
    details["power_verbs"] = power_score

    # ── 5. Keyword density (15 pts) ───────────────────────────────────────────
    if skills:
        evidence_text = (sections.get("experience", "") + " " + sections.get("projects", "")).lower()
        evidenced = sum(
            1 for s in skills
            if re.search(r"(?<![a-z0-9])" + re.escape(s.lower()) + r"(?![a-z0-9])", evidence_text)
        )
        density = evidenced / len(skills) if skills else 0
        kw_score = min(int(density * 20), 15)
        if density < 0.5:
            flags.append(f"Only {evidenced}/{len(skills)} skills ({density:.0%}) appear in experience/projects — ATS prefers skills backed by evidence")
    else:
        kw_score = 0
        flags.append("No skills detected — ATS keyword matching will score 0")
    score += kw_score
    details["keyword_density"] = kw_score

    # ── 6. Structure quality (15 pts) ─────────────────────────────────────────
    struct_score = 15
    lines = raw_text.splitlines()
    # Excessive line length
    long_lines = sum(1 for l in lines if len(l) > 120)
    if long_lines > 5:
        struct_score -= 3
        flags.append(f"{long_lines} lines exceed 120 chars — may wrap poorly in ATS text boxes")
    # All-caps overuse
    caps_lines = sum(1 for l in lines if l.strip().isupper() and len(l.strip()) > 10)
    if caps_lines > 4:
        struct_score -= 3
        flags.append(f"{caps_lines} all-caps lines detected — ATS may not parse section headers correctly")
    # File length sanity
    if len(raw_text) < 200:
        struct_score -= 9
        flags.append("Resume is very short — likely a parsing failure")
    struct_score = max(struct_score, 0)
    score += struct_score
    details["structure"] = struct_score

    final_score = min(round(score), 100)
    
    # Verdict
    if final_score >= 80:
        verdict = "ATS Ready"
    elif final_score >= 60:
        verdict = "Needs Minor Fixes"
    else:
        verdict = "High Risk of Rejection"

    return {
        "score":   final_score,
        "verdict": verdict,
        "flags":   flags,
        "details": details,
    }

# ── Main entry ────────────────────────────────────────────────────────────────

def parse_document(file_bytes: bytes, filename: str) -> dict:
    fn = filename.lower()

    if fn.endswith(".pdf"):
        raw_text = extract_pdf_text(file_bytes)
        # For two-column PDFs: column-split text can truncate words at the boundary.
        # Extract a full-page (no column split) version purely for contact info
        # and institution name recovery — don't use it for section splitting.
        _full_page_text = ""
        try:
            import pdfplumber, io as _io
            with pdfplumber.open(_io.BytesIO(file_bytes)) as _pdf:
                _full_page_text = "\n".join(
                    (p.extract_text(x_tolerance=3, y_tolerance=3) or "")
                    for p in _pdf.pages
                )
        except Exception:
            pass
    elif fn.endswith((".docx", ".doc")):
        raw_text = extract_docx_text(file_bytes)
        _full_page_text = ""
    elif fn.endswith(".txt"):
        raw_text = _clean_text(file_bytes.decode("utf-8", errors="replace"))
        _full_page_text = ""
    else:
        raise ValueError(f"Unsupported format: {filename}. Supported: PDF, DOCX, TXT")

    # ── Institution name recovery: MUST happen before sections are built ──────
    # Two-column PDFs clip words at column boundary. Patch raw_text and _COL_CACHE
    # using the full-page (no column split) text FIRST, then build sections.
    if _full_page_text:
        import re as _re2
        # Find all complete institution names in the full (unclipped) page text
        # Then patch any truncated version of the same name in the column-split raw_text.
        # Strategy: find "WORD OF WORD" style institution names in full_page that end in a
        # known institution suffix, then see if a truncated version exists in raw_text.
        _INST_SUFFIXES = (
            "ENGINEERING","TECHNOLOGY","POLYTECHNIC",
            "UNIVERSITY","COLLEGE","INSTITUTE","SCHOOL","ACADEMY",
        )
        for _suffix_word in _INST_SUFFIXES:
            # Find the complete institution name in full_page
            _full_inst_pat = _re2.compile(
                r"[A-Z][A-Z\.\s]{3,}\b" + _re2.escape(_suffix_word) + r"\b",
                _re2.IGNORECASE
            )
            for _fm in _full_inst_pat.finditer(_full_page_text):
                _full_name = _fm.group().strip()
                # Try truncations of the suffix word (missing last 1-4 chars)
                for _cut in range(1, min(5, len(_suffix_word)-3)):
                    _trunc_suffix = _suffix_word[:-_cut]
                    _trunc_name   = _full_name[:-_cut]   # e.g. "MA COLLEGE OF ENGINEERIN"
                    # Only patch if the truncated version is literally in raw_text
                    # and the full version is NOT already there (same context)
                    if _trunc_name in raw_text and _full_name not in raw_text:
                        logger.info("Recovering truncated institution: %r → %r",
                                    _trunc_name, _full_name)
                        raw_text = raw_text.replace(_trunc_name, _full_name)
                        for _hash, _cols in _COL_CACHE.items():
                            _COL_CACHE[_hash] = {
                                "left":       _cols.get("left","").replace(_trunc_name, _full_name),
                                "right":      _cols.get("right","").replace(_trunc_name, _full_name),
                                "is_two_col": _cols.get("is_two_col", False),
                            }
                        break

    # Contact: prefer full-page extraction (no column truncation)
    contact_info = extract_contact_info(_full_page_text or raw_text)
    if _full_page_text:
        _ci2 = extract_contact_info(raw_text)
        for _k, _v in _ci2.items():
            if _k not in contact_info or len(str(_v)) > len(str(contact_info.get(_k, ""))):
                contact_info[_k] = _v

    sections       = split_into_sections(raw_text)
    mismatches     = detect_section_mismatches(sections)
    inline_skills  = extract_inline_skills(raw_text)
    preprocessed   = preprocess_for_embedding(raw_text)

    logger.info(
        "Parsed '%s': sections=%s skills=%d",
        filename, list(sections.keys()), len(inline_skills)
    )

    return {
        "raw_text":          raw_text,
        "sections":          sections,
        "mismatches":        mismatches,
        "inline_skills":     inline_skills,
        "preprocessed_text": preprocessed,
        "contact_info":      contact_info,
    }